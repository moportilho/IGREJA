import streamlit as st
import pandas as pd
import time
import datetime
import pyodbc
from io import BytesIO

# Bibliotecas necessárias para geração de documentos
from fpdf import FPDF
from docx import Document

# -----------------------------------------------------------------------------
# Configuração Geral do Streamlit
# -----------------------------------------------------------------------------
st.set_page_config(page_title="Demonstração Local", layout="wide")

# -----------------------------------------------------------------------------
# Helpers de data (somente EXIBIÇÃO em PT-BR)
# -----------------------------------------------------------------------------
def fmt_date_br(d):
    """Converte uma data para 'DD/MM/YYYY' somente para exibição."""
    if d is None or d == "":
        return ""
    if isinstance(d, (datetime.date, datetime.datetime)):
        return d.strftime("%d/%m/%Y")
    # tenta normalizar com pandas
    ts = pd.to_datetime(d, errors="coerce")
    if pd.isna(ts):
        return ""
    return ts.strftime("%d/%m/%Y")

def df_to_br_display(df, date_cols):
    """Retorna uma cópia do DF com colunas de data formatadas como 'DD/MM/YYYY'."""
    df2 = df.copy()
    for c in date_cols:
        if c in df2.columns:
            s = pd.to_datetime(df2[c], errors="coerce")
            s = s.dt.strftime("%d/%m/%Y")
            df2[c] = s.fillna("")
    return df2

def safe_date(v, default=None):
    if isinstance(v, datetime.datetime):
        return v.date()
    if isinstance(v, datetime.date):
        return v
    return default or datetime.date.today()

# -----------------------------------------------------------------------------
# Conexão com banco de dados
# -----------------------------------------------------------------------------

def get_connection():
    server = st.secrets["server"]
    database = st.secrets["database"]
    username = st.secrets["username"]
    password = st.secrets["password"]
    driver = '{ODBC Driver 17 for SQL Server}'
    try:
        conx = pyodbc.connect(
            'Driver='+ driver + ';Server='+ server + ';Database=' + database + ';Uid=' + username + ';Pwd={' + password + '}'
        )
        return conx
    except Exception as e:
        st.error(f"Erro ao conectar ao banco de dados: {e}")
        return None


def read_records(query, params=None):
    """
    Executa um SELECT e retorna um DataFrame.
    """
    conx = get_connection()
    if not conx:
        return pd.DataFrame()
    try:
        df = pd.read_sql(query, conx, params=params)
        return df
    except Exception as e:
        st.error(f"Erro ao ler registros: {e}")
        return pd.DataFrame()
    finally:
        conx.close()


# -----------------------------------------------------------------------------
# Executa um comando SQL (INSERT, UPDATE ou DELETE) e confirma (commit)
# -----------------------------------------------------------------------------

def execute_query(query, params=None):
    """
    Executa um comando SQL (INSERT, UPDATE ou DELETE) e confirma (commit).
    """
    conx = get_connection()
    if not conx:
        return False
    try:
        cursor = conx.cursor()
        if params:
            cursor.execute(query, params)
        else:
            cursor.execute(query)
        conx.commit()
        if cursor.rowcount == 0:
            return "Nenhuma linha foi afetada."
        return True

    except Exception as e:
        error_message = str(e)
        # Se achar "2627" ou "duplicate key", é chave duplicada
        if "2627" in error_message or "duplicate key" in error_message:
            return "Entrada duplicada."  # Mensagem genérica
        else:
            return error_message
    finally:
        conx.close()


# -----------------------------------------------------------------------------
# SEÇÃO DE LOGIN
# -----------------------------------------------------------------------------

CREDENTIALS = {
    "adm": "adm123",
    "adm-financeiro": "fin123",
    "adm-secretaria": "sec123"
}

if "logged_in" not in st.session_state:
    st.session_state["logged_in"] = False
if "user_role" not in st.session_state:
    st.session_state["user_role"] = None


def login_screen():
    st.title("Login do Sistema")
    user = st.text_input("Usuário")
    password = st.text_input("Senha", type="password")
    if st.button("Entrar"):
        if user in CREDENTIALS and password == CREDENTIALS[user]:
            st.session_state["logged_in"] = True
            st.session_state["user_role"] = user
            st.rerun()
        else:
            st.error("Usuário ou senha inválidos. Tente novamente.")


def logout_button():
    if st.sidebar.button("Sair"):
        st.session_state["logged_in"] = False
        st.session_state["user_role"] = None
        st.rerun()


# -----------------------------------------------------------------------------
# INICIALIZAÇÃO DE DADOS (session_state)
# -----------------------------------------------------------------------------

# 1) Igreja
if "igreja_data" not in st.session_state:
    st.session_state["igreja_data"] = {
        "logotipo": None,
        "cnpj": "",
        "data_abertura": None,
        "endereco": "",
        "pastor_nome": "",
        "pastor_entrada": None,
        "pastor_saida": None
    }

# 2) Membros (ajustado para novo esquema: PK 'id' e nova coluna 'matricula')
if "membros_data" not in st.session_state:
    st.session_state["membros_data"] = pd.DataFrame(columns=[
        "id", "matricula", "nome", "foto", "ministerio", "endereco",
        "telefone", "email", "sexo", "data_nascimento", "estado_civil", "nome_conjuge",
        "disciplina_data_ini", "disciplina_data_fim", "data_entrada",
        "tipo_entrada", "data_desligamento", "motivo_desligamento",
        "mes_aniversario"
    ])


# -----------------------------------------------------------------------------
# PÁGINA 1: Cadastro da Igreja
# -----------------------------------------------------------------------------

def page_igreja():
    st.header("Cadastro de Igreja")

    # Verifica se já existe um cadastro na tabela Igreja
    df_igreja = read_records("SELECT * FROM Igreja")
    
    if st.session_state["user_role"] == "adm-secretaria":
        st.subheader("Igreja Cadastrada")
        if df_igreja.empty:
            st.info("Nenhuma igreja cadastrada.")
        else:
            igreja = df_igreja.iloc[0]
            if igreja.get("logotipo") is not None:
                st.image(igreja["logotipo"], caption="Logotipo da Igreja", width=200)
            st.write(f"**CNPJ:** {igreja['cnpj']}")
            st.write(f"**Data de Abertura:** {fmt_date_br(igreja['data_abertura'])}")
            st.write(f"**Endereço:** {igreja['endereco']}")
            st.write(f"**Nome do Pastor:** {igreja['pastor_nome']}")
            st.write(f"**Data de Entrada do Pastor:** {fmt_date_br(igreja['pastor_entrada'])}")
            st.write(f"**Data de Saída do Pastor:** {fmt_date_br(igreja['pastor_saida'])}")
        return

    # Para usuários com permissões de edição (adm, adm-financeiro)
    if not df_igreja.empty:
        # Considerando que há apenas um registro
        igreja = df_igreja.iloc[0]
        
        # Inicializa a flag de edição, se não existir
        if "editar_igreja" not in st.session_state:
            st.session_state["editar_igreja"] = False

        # Modo de Edição
        if st.session_state["editar_igreja"]:
            st.subheader("Editar Dados da Igreja")
            with st.form("form_editar_igreja"):
                uploaded_logo = st.file_uploader("Selecione um novo logotipo da Igreja (opcional)", 
                                                   type=["png", "jpg", "jpeg"])
                if uploaded_logo is not None:
                    logotipo_bin = uploaded_logo.read()
                else:
                    logotipo_bin = igreja["logotipo"]
                
                cnpj_val = st.text_input("CNPJ da Igreja*", value=igreja["cnpj"], disabled=True)

                try:
                    data_abertura_val = st.date_input("Data de Abertura*", value=safe_date(igreja["data_abertura"]), format="DD/MM/YYYY")
                except TypeError:
                    data_abertura_val = st.date_input("Data de Abertura*", value=safe_date(igreja["data_abertura"]))

                endereco_val = st.text_input("Endereço*", value=igreja["endereco"])
                pastor_nome_val = st.text_input("Nome do Pastor*", value=igreja["pastor_nome"])
                
                col1, col2 = st.columns(2)
                with col1:
                    try:
                        pastor_entrada_val = st.date_input("Data de Entrada do Pastor*", value=safe_date(igreja["pastor_entrada"]), format="DD/MM/YYYY")
                    except TypeError:
                        pastor_entrada_val = st.date_input("Data de Entrada do Pastor*", value=safe_date(igreja["pastor_entrada"]))
                with col2:
                    try:
                        pastor_saida_val = st.date_input("Data de Saída do Pastor*", value=safe_date(igreja["pastor_saida"]), format="DD/MM/YYYY")
                    except TypeError:
                        pastor_saida_val = st.date_input("Data de Saída do Pastor*", value=safe_date(igreja["pastor_saida"]))
                
                submit_edit = st.form_submit_button("Atualizar Dados")
                
                if submit_edit:
                    update_sql = """
                        UPDATE Igreja
                        SET logotipo = ?,
                            data_abertura = ?,
                            endereco = ?,
                            pastor_nome = ?,
                            pastor_entrada = ?,
                            pastor_saida = ?
                        WHERE cnpj = ?
                    """
                    params_update = (
                        logotipo_bin,
                        data_abertura_val,
                        endereco_val,
                        pastor_nome_val,
                        pastor_entrada_val,
                        pastor_saida_val,
                        igreja["cnpj"]
                    )
                    ok_update = execute_query(update_sql, params_update)
                    if ok_update:
                        st.success("Dados da Igreja atualizados com sucesso!")
                        st.session_state["editar_igreja"] = False
                        st.rerun()
                    else:
                        st.error("Falha ao atualizar os dados da Igreja.")
            
            if st.button("Cancelar Edição"):
                st.session_state["editar_igreja"] = False
                st.rerun()
                
        # Modo de Visualização
        else:
            st.subheader("Igreja Cadastrada")
            if igreja.get("logotipo") is not None:
                st.image(igreja["logotipo"], caption="Logotipo da Igreja", width=200)
            st.write(f"**CNPJ:** {igreja['cnpj']}")
            st.write(f"**Data de Abertura:** {fmt_date_br(igreja['data_abertura'])}")
            st.write(f"**Endereço:** {igreja['endereco']}")
            st.write(f"**Nome do Pastor:** {igreja['pastor_nome']}")
            st.write(f"**Data de Entrada do Pastor:** {fmt_date_br(igreja['pastor_entrada'])}")
            st.write(f"**Data de Saída do Pastor:** {fmt_date_br(igreja['pastor_saida'])}")
            
            col1, col2 = st.columns(2)
            with col1:
                if st.button("Editar Dados da Igreja"):
                    st.session_state["editar_igreja"] = True
                    st.rerun()
            with col2:
                if st.button("Excluir Igreja"):
                    delete_sql = "DELETE FROM Igreja WHERE cnpj = ?"
                    ok_delete = execute_query(delete_sql, (igreja["cnpj"],))
                    if ok_delete:
                        st.success("Igreja excluída com sucesso!")
                        st.rerun()
                    else:
                        st.error("Falha ao excluir a Igreja.")
                        
    else:
        # Caso não exista nenhum cadastro, exibe o formulário de cadastro
        st.subheader("Cadastrar Nova Igreja")
        with st.form("form_nova_igreja"):
            uploaded_logo = st.file_uploader("Selecione o logotipo da Igreja (opcional)", 
                                             type=["png", "jpg", "jpeg"])
            logotipo_bin = uploaded_logo.read() if uploaded_logo is not None else None

            cnpj_val = st.text_input("CNPJ da Igreja*")
            try:
                data_abertura_val = st.date_input("Data de Abertura*", value=None, min_value=datetime.date(1900, 1, 1), format="DD/MM/YYYY")
            except TypeError:
                data_abertura_val = st.date_input("Data de Abertura*", value=None, min_value=datetime.date(1900, 1, 1))
            endereco_val = st.text_input("Endereço*")
            pastor_nome_val = st.text_input("Nome do Pastor*")

            col1, col2 = st.columns(2)
            with col1:
                try:
                    pastor_entrada_val = st.date_input("Data de Entrada do Pastor*", value=None, min_value=datetime.date(1900, 1, 1), format="DD/MM/YYYY")
                except TypeError:
                    pastor_entrada_val = st.date_input("Data de Entrada do Pastor*", value=None, min_value=datetime.date(1900, 1, 1))
            with col2:
                try:
                    pastor_saida_val = st.date_input("Data de Saída do Pastor*", value=None, min_value=datetime.date(1900, 1, 1), format="DD/MM/YYYY")
                except TypeError:
                    pastor_saida_val = st.date_input("Data de Saída do Pastor*", value=None, min_value=datetime.date(1900, 1, 1))

            submit_button = st.form_submit_button("Salvar Dados da Igreja")
            if submit_button:
                erros = []
                if not cnpj_val.strip():
                    erros.append("CNPJ da Igreja")
                if not endereco_val.strip():
                    erros.append("Endereço")
                if not pastor_nome_val.strip():
                    erros.append("Nome do Pastor")
                if not data_abertura_val:
                    erros.append("Data de Abertura")
                if not pastor_entrada_val:
                    erros.append("Data de Entrada do Pastor")
                if not pastor_saida_val:
                    erros.append("Data de Saída do Pastor")
                if erros:
                    texto_erros = "### Atenção!\n\nOs seguintes campos obrigatórios não foram preenchidos:\n\n"
                    for campo in erros:
                        texto_erros += f"- {campo}\n"
                    st.error(texto_erros)
                else:
                    insert_sql = """
                        INSERT INTO Igreja (
                            logotipo,
                            cnpj,
                            data_abertura,
                            endereco,
                            pastor_nome,
                            pastor_entrada,
                            pastor_saida
                        )
                        VALUES (?, ?, ?, ?, ?, ?, ?)
                    """
                    params_insert = (
                        logotipo_bin,
                        cnpj_val,
                        data_abertura_val,
                        endereco_val,
                        pastor_nome_val,
                        pastor_entrada_val,
                        pastor_saida_val
                    )
                    ok = execute_query(insert_sql, params_insert)
                    if ok:
                        st.success("Dados da Igreja salvos com sucesso!")
                        st.rerun()
                    else:
                        st.error("Falha ao inserir dados da Igreja.")

# -----------------------------------------------------------------------------
# PÁGINA 2: Cadastro de Membros
# -----------------------------------------------------------------------------

def page_membros():
    st.header("Cadastro de Membros")
    df_membros = read_records("SELECT * FROM Membros")

    # Para secretaria: apenas visualização
    if st.session_state["user_role"] == "adm-secretaria":
        if df_membros.empty:
            st.info("Ainda não há nenhum membro adicionado.")
        else:
            st.subheader("Listagem de Membros")
            cols_dt_m = ["data_nascimento","disciplina_data_ini","disciplina_data_fim","data_entrada","data_desligamento"]
            df_membros_br = df_to_br_display(df_membros, cols_dt_m)
            st.dataframe(df_membros_br, use_container_width=True)
            st.subheader("Fotos dos Membros")
            for _, row in df_membros.iterrows():
                if row.get("foto") is not None:
                    st.image(row["foto"], caption=row['nome'], width=100)
        return

    if df_membros.empty:
        st.info("Ainda não há nenhum membro adicionado.")

    # Sincroniza com session_state
    st.session_state["membros_data"] = df_membros
    membros_df = st.session_state["membros_data"]

    # Formulário de adicionar novo membro
    with st.expander("Adicionar Membro"):
        with st.form("form_add_membro"):
            matricula_input = st.text_input("Matrícula*")
            nome = st.text_input("Nome completo*")
            uploaded_foto = st.file_uploader("Foto do Membro (opcional)", type=["png", "jpg", "jpeg"])
            foto_bytes = uploaded_foto.read() if uploaded_foto is not None else None
            ministerio = st.text_input("Ministério (opcional)")
            endereco = st.text_input("Endereço (opcional)")
            telefone = st.text_input("Telefone (opcional) (DDXXXXXXXXX)")
            email = st.text_input("E-mail (opcional)")
            sexo = st.selectbox("Sexo*", ["Selecione...", "Masculino", "Feminino", "Outro"])

            try:
                data_nascimento = st.date_input("Data de Nascimento*", min_value=datetime.date(1900, 1, 1), format="DD/MM/YYYY")
            except TypeError:
                data_nascimento = st.date_input("Data de Nascimento*", min_value=datetime.date(1900, 1, 1))
            if data_nascimento:
                st.caption(f"Selecionado: {data_nascimento.strftime('%d/%m/%Y')}")

            estado_civil = st.selectbox("Estado Civil (opcional)", ["Selecione...", "Solteiro(a)", "Casado(a)", "Divorciado(a)", "Viúvo(a)"])
            nome_conjuge = st.text_input("Nome do Cônjuge (se Casado)")

            try:
                disciplina_data_ini = st.date_input("Data de início Disciplina (opcional)", min_value=datetime.date(1900, 1, 1), format="DD/MM/YYYY")
            except TypeError:
                disciplina_data_ini = st.date_input("Data de início Disciplina (opcional)", min_value=datetime.date(1900, 1, 1))

            try:
                disciplina_data_fim = st.date_input("Data de saída Disciplina (opcional)", min_value=datetime.date(1900, 1, 1), format="DD/MM/YYYY")
            except TypeError:
                disciplina_data_fim = st.date_input("Data de saída Disciplina (opcional)", min_value=datetime.date(1900, 1, 1))

            try:
                data_entrada = st.date_input("Data de entrada (Ativo) (opcional)", value=datetime.date.today(), min_value=datetime.date(1900, 1, 1), format="DD/MM/YYYY")
            except TypeError:
                data_entrada = st.date_input("Data de entrada (Ativo) (opcional)", value=datetime.date.today(), min_value=datetime.date(1900, 1, 1))

            tipo_entrada = st.selectbox("Tipo de entrada*", ["Selecione...", "Batismo", "Transferência", "Aclamação", "Reconciliação"])

            try:
                data_desligamento = st.date_input("Data do desligamento (Inativo) (opcional)", min_value=datetime.date(1900, 1, 1), format="DD/MM/YYYY")
            except TypeError:
                data_desligamento = st.date_input("Data do desligamento (Inativo) (opcional)", min_value=datetime.date(1900, 1, 1))

            motivo_desligamento = st.selectbox("Motivo do Desligamento (opcional)", ["Nenhum", "A pedido", "Ausência", "Transferência", "Outra denominação", "Outros motivos"])

            submitted = st.form_submit_button("Adicionar Membro")
            if submitted:
                erros = []
                if not matricula_input.strip():
                    erros.append("Matrícula")
                else:
                    try:
                        matricula_val = int(matricula_input)
                    except ValueError:
                        erros.append("Matrícula deve ser um número inteiro")
                if not nome.strip():
                    erros.append("Nome completo")
                if not data_nascimento:
                    erros.append("Data de Nascimento")
                if sexo == "Selecione...":
                    erros.append("Sexo")

                if erros:
                    st.error("Preencha corretamente os campos obrigatórios: " + ", ".join(erros))
                else:
                    mes_aniversario = data_nascimento.month if data_nascimento else None

                    insert_sql = """
                        INSERT INTO Membros (
                            matricula, nome, foto, ministerio, endereco, telefone, email, sexo,
                            data_nascimento, estado_civil, nome_conjuge,
                            disciplina_data_ini, disciplina_data_fim, data_entrada,
                            tipo_entrada, data_desligamento, motivo_desligamento,
                            mes_aniversario
                        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """
                    params = (
                        matricula_val,
                        nome,
                        foto_bytes,
                        ministerio if ministerio.strip() else None,
                        endereco if endereco.strip() else None,
                        telefone if telefone.strip() else None,
                        email if email.strip() else None,
                        sexo,
                        data_nascimento,  # armazenado como DATE
                        None if estado_civil == "Selecione..." else estado_civil,
                        nome_conjuge if nome_conjuge.strip() else None,
                        disciplina_data_ini,
                        disciplina_data_fim,
                        data_entrada,
                        None if tipo_entrada == "Selecione..." else tipo_entrada,
                        data_desligamento,
                        motivo_desligamento,
                        mes_aniversario
                    )
                    sucesso = execute_query(insert_sql, params)
                    if sucesso is True:
                        st.success("Membro inserido com sucesso!")
                        st.rerun()
                    else:
                        st.error(f"Falha ao inserir membro: {sucesso}")

    # Exibição e edição dos membros
    st.subheader("Listagem de Membros")

    df_editor = membros_df.drop(columns=["foto"], errors='ignore')
    edited_df = st.data_editor(
        df_editor,
        hide_index=True,
        use_container_width=True,
        key="membros_editor",
        column_config={
            "id": st.column_config.TextColumn("ID", disabled=True),
            "data_nascimento": st.column_config.DateColumn("Data de Nascimento", format="DD/MM/YYYY"),
            "disciplina_data_ini": st.column_config.DateColumn("Início Disciplina", format="DD/MM/YYYY"),
            "disciplina_data_fim": st.column_config.DateColumn("Fim Disciplina", format="DD/MM/YYYY"),
            "data_entrada": st.column_config.DateColumn("Data de Entrada", format="DD/MM/YYYY"),
            "data_desligamento": st.column_config.DateColumn("Data de Desligamento", format="DD/MM/YYYY"),
        }
    )

    # -----------------------------------------
    # Atualizar / Remover foto do membro
    # -----------------------------------------
    with st.expander("🖼️ Atualizar foto do membro"):
        if membros_df.empty:
            st.info("Cadastre membros para poder editar a foto.")
        else:
            membro_sel = st.selectbox(
                "Selecione o membro",
                options=membros_df["id"],
                format_func=lambda i: membros_df.loc[membros_df["id"] == i, "nome"].values[0]
            )

            # Mostra a foto atual (se existir)
            row_atual = membros_df.loc[membros_df["id"] == membro_sel].iloc[0]
            st.write(f"**Membro:** {row_atual['nome']}")
            if row_atual.get("foto") is not None:
                st.image(row_atual["foto"], caption="Foto atual", width=150)
            else:
                st.caption("Sem foto cadastrada.")

            col_f1, col_f2 = st.columns(2)
            with col_f1:
                nova_foto = st.file_uploader(
                    "Nova foto (PNG/JPG/JPEG)",
                    type=["png", "jpg", "jpeg"],
                    key="upload_foto_edit"
                )
                if st.button("Salvar nova foto"):
                    if nova_foto is None:
                        st.warning("Envie uma imagem antes de salvar.")
                    else:
                        foto_bytes_nova = nova_foto.read()
                        ok = execute_query(
                            "UPDATE Membros SET foto = ? WHERE id = ?",
                            (foto_bytes_nova, int(membro_sel))
                        )
                        if ok is True:
                            st.success("Foto atualizada com sucesso!")
                            st.rerun()
                        else:
                            st.error(f"Falha ao atualizar foto: {ok}")

            with col_f2:
                if st.button("Remover foto atual"):
                    ok = execute_query(
                        "UPDATE Membros SET foto = NULL WHERE id = ?",
                        (int(membro_sel),)
                    )
                    if ok is True:
                        st.success("Foto removida com sucesso!")
                        st.rerun()
                    else:
                        st.error(f"Falha ao remover foto: {ok}")


    # Salvar edições no banco
    if st.button("Salvar Alterações de Edição"):
        for _, row in edited_df.iterrows():
            update_sql = """
                UPDATE Membros
                SET matricula = ?, nome = ?, ministerio = ?, endereco = ?, telefone = ?, email = ?, sexo = ?,
                    data_nascimento = ?, estado_civil = ?, nome_conjuge = ?,
                    disciplina_data_ini = ?, disciplina_data_fim = ?, data_entrada = ?,
                    tipo_entrada = ?, data_desligamento = ?, motivo_desligamento = ?, mes_aniversario = ?
                WHERE id = ?
            """
            update_params = (
                row.get("matricula"),
                row.get("nome"),
                row.get("ministerio"),
                row.get("endereco"),
                row.get("telefone"),
                row.get("email"),
                row.get("sexo"),
                row.get("data_nascimento"),
                row.get("estado_civil"),
                row.get("nome_conjuge"),
                row.get("disciplina_data_ini"),
                row.get("disciplina_data_fim"),
                row.get("data_entrada"),
                row.get("tipo_entrada"),
                row.get("data_desligamento"),
                row.get("motivo_desligamento"),
                row.get("mes_aniversario"),
                row.get("id")
            )
            execute_query(update_sql, update_params)
        st.success("Alterações atualizadas!")
        st.rerun()

    # Exclusão de membro
    with st.expander("Excluir Membro"):
        if "id" in membros_df.columns:
            lista_ids = list(membros_df["id"].dropna().unique())
        else:
            lista_ids = []
        if not lista_ids:
            st.info("Não há membros cadastrados para exclusão.")
        else:
            id_excluir = st.selectbox(
                "ID do membro a excluir",
                options=lista_ids
            )
            if st.button("Confirmar Exclusão"):
                try:
                    id_param = int(id_excluir)
                except (ValueError, TypeError):
                    st.error("ID inválido.")
                    return
                delete_sql = "DELETE FROM Membros WHERE id = ?"
                sucesso = execute_query(delete_sql, (id_param,))
                if sucesso is True:
                    st.success(f"Membro de ID {id_param} excluído.")
                    st.session_state["membros_data"] = read_records("SELECT * FROM Membros")
                    st.rerun()
                else:
                    st.error(f"Falha ao excluir membro: {sucesso}")

    # Pré-visualizar fotos
    st.subheader("Fotos dos Membros")
    for _, row in st.session_state["membros_data"].iterrows():
        if row.get("foto") is not None:
            st.image(row["foto"], caption=row['nome'], width=100)

# -----------------------------------------------------------------------------
# PÁGINA 3: Relatórios
# -----------------------------------------------------------------------------

def page_relatorios():
    st.header("Relatórios")

    membros_df = st.session_state["membros_data"]

    col1, col2, col3 = st.columns(3)

    # 1) Geração de PDF do Certificado de Batismo
    with col1:
        if st.button("Gerar PDF - Certificado de Batismo"):
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=16, style='B')
            pdf.cell(200, 10, txt="CERTIFICADO DE BATISMO", ln=1, align='C')
            pdf.set_font("Arial", size=12)

            pdf.ln(10)
            pdf.multi_cell(0, 10, txt=(
                "Declaramos que o membro [NOME] recebeu o Santo Batismo nesta igreja,\n"
                "conforme as doutrinas cristãs, no dia [DATA].\n\n"
                "Assinatura:\n"
                "_________________________________________"
            ))

            pdf_data = pdf.output(dest="S").encode("latin-1")
            st.download_button(
                label="Baixar PDF Certificado",
                data=pdf_data,
                file_name="certificado_batismo.pdf",
                mime="application/pdf"
            )

    # 2) Geração de Word (DOCX) da Carta de Transferência
    with col2:
        if st.button("Gerar Word - Carta de Transferência"):
            doc = Document()
            doc.add_heading("CARTA DE TRANSFERÊNCIA", 0)

            p = doc.add_paragraph()
            p.add_run("Aos cuidados da Igreja de destino,\n\n").bold = True
            p.add_run(
                "Certificamos que o(a) membro [NOME] faz parte de nossa congregação, "
                "estando em comunhão, e solicitou transferência para a Igreja [DESTINO]. "
                "Concede-se, portanto, esta carta para os devidos fins.\n\n"
            )
            p.add_run("Atenciosamente,\n[Igreja de Origem]")

            doc_buffer = BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)

            st.download_button(
                label="Baixar Word Transferência",
                data=doc_buffer.getvalue(),
                file_name="carta_transferencia.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    # 3) Geração de PDF - Carta por Ausência
    with col3:
        if st.button("Gerar PDF - Carta por Ausência"):
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=16, style='B')
            pdf.cell(200, 10, txt="CARTA POR AUSÊNCIA", ln=1, align='C')
            pdf.set_font("Arial", size=12)

            pdf.ln(10)
            pdf.multi_cell(0, 10, txt=(
                "Ao(À) Sr(a). [NOME DO MEMBRO],\n\n"
                "Consta em nossos registros que o(a) senhor(a) se encontra ausente de nossas atividades "
                "e cultos por período prolongado. Solicitamos o comparecimento ou contato para "
                "regularização de seu estado como membro ativo.\n\n"
                "Atenciosamente,\n[Igreja]"
            ))

            pdf_data = pdf.output(dest="S").encode("latin-1")
            st.download_button(
                label="Baixar PDF Carta Ausência",
                data=pdf_data,
                file_name="carta_ausencia.pdf",
                mime="application/pdf"
            )

    # 4) Geração de Excel com os membros cadastrados
    st.subheader("Gerar Excel dos Membros Cadastrados")
    if st.button("Gerar Excel"):
        output = BytesIO()
        membros_sem_foto = st.session_state["membros_data"].drop(columns=["foto"], errors="ignore")
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            membros_sem_foto.to_excel(writer, sheet_name="Membros", index=False)
        st.download_button(
            label="Baixar Excel com Membros",
            data=output.getvalue(),
            file_name="relatorio_membros.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )


# -----------------------------------------------------------------------------
# Página 4 (exclusiva para adm-financeiro): Página Financeira
# -----------------------------------------------------------------------------
def ensure_finance_schema():
    """
    Garante a existência da tabela/índices financeiros no SQL Server.
    """
    ddl = """
    IF NOT EXISTS (SELECT * FROM sys.objects WHERE object_id = OBJECT_ID(N'[dbo].[DizimoLancamentos]') AND type in (N'U'))
    BEGIN
        CREATE TABLE [dbo].[DizimoLancamentos](
            [id]                INT IDENTITY(1,1) PRIMARY KEY,
            [membro_id]         INT NOT NULL,
            [competencia]       DATE NOT NULL,
            [valor_dizimo]      DECIMAL(10,2) NOT NULL DEFAULT 0,
            [valor_oferta]      DECIMAL(10,2) NOT NULL DEFAULT 0,
            [data_pagamento]    DATE NOT NULL,
            [forma_pagamento]   VARCHAR(30) NULL,
            [observacoes]       NVARCHAR(255) NULL,
            [criado_em]         DATETIME2 NOT NULL DEFAULT SYSUTCDATETIME(),
            [atualizado_em]     DATETIME2 NULL
        );
        ALTER TABLE [dbo].[DizimoLancamentos]
            ADD CONSTRAINT FK_Dizimos_Membro
            FOREIGN KEY ([membro_id]) REFERENCES [dbo].[Membros]([id]) ON DELETE CASCADE;

        ALTER TABLE [dbo].[DizimoLancamentos]
            ADD [ano] AS (YEAR([competencia])) PERSISTED,
                [mes] AS (MONTH([competencia])) PERSISTED;

        CREATE UNIQUE INDEX UX_Dizimo_MembroCompetencia
            ON [dbo].[DizimoLancamentos]([membro_id], [ano], [mes]);
    END
    """
    _ = execute_query(ddl)  # Ignora retorno; se já existir, não faz nada

def page_financeiro():
    ensure_finance_schema()  # garante tabela/índices

    st.header("Página Financeira • Dízimos e Ofertas")

    # ==== TABS ====
    tab1, tab2, tab3 = st.tabs(["➕ Lançar contribuição", "📊 Painel anual (estilo planilha)", "🧾 Gerenciar lançamentos"])

    # ========= TAB 1: Lançar contribuição =========
    with tab1:
        st.subheader("Registrar contribuição mensal")

        membros = read_records("SELECT id, nome FROM Membros ORDER BY nome")
        if membros.empty:
            st.info("Cadastre membros antes de lançar contribuições.")
        else:
            colA, colB, colC = st.columns(3)
            with colA:
                membro_escolhido = st.selectbox("Membro*", options=membros["id"], format_func=lambda i: membros.loc[membros["id"]==i, "nome"].values[0])
            with colB:
                ano = st.number_input("Ano (competência)*", min_value=1900, max_value=2100, value=datetime.date.today().year, step=1)
            with colC:
                mes = st.selectbox("Mês (competência)*", options=list(range(1,13)), format_func=lambda m: datetime.date(2000, m, 1).strftime("%B").capitalize())

            # 1º dia do mês como competência (ex.: 2025-08-01)
            competencia = datetime.date(int(ano), int(mes), 1)

            col1, col2, col3 = st.columns(3)
            with col1:
                valor_dizimo = st.number_input("Valor do dízimo (R$)*", min_value=0.0, step=10.0, format="%.2f")
            with col2:
                valor_oferta = st.number_input("Valor de oferta (R$)", min_value=0.0, step=5.0, value=0.0, format="%.2f")
            with col3:
                try:
                    data_pagamento = st.date_input("Data do pagamento*", value=datetime.date.today(), min_value=datetime.date(1900,1,1), format="DD/MM/YYYY")
                except TypeError:
                    data_pagamento = st.date_input("Data do pagamento*", value=datetime.date.today(), min_value=datetime.date(1900,1,1))

            col4, col5 = st.columns(2)
            with col4:
                forma_pagamento = st.selectbox("Forma de pagamento", ["Dinheiro", "Pix", "Cartão", "Transferência", "Boleto", "Outro"])
            with col5:
                observacoes = st.text_input("Observações (opcional)")

            if st.button("Salvar / Atualizar"):
                # INSERT; se já existir para (membro, ano, mes), faz UPDATE
                insert_sql = """
                    INSERT INTO DizimoLancamentos (membro_id, competencia, valor_dizimo, valor_oferta, data_pagamento, forma_pagamento, observacoes)
                    VALUES (?, ?, ?, ?, ?, ?, ?)
                """
                params = (int(membro_escolhido), competencia, float(valor_dizimo), float(valor_oferta), data_pagamento, forma_pagamento, observacoes if observacoes.strip() else None)
                result = execute_query(insert_sql, params)

                if result is True:
                    st.success("Lançamento salvo.")
                elif isinstance(result, str) and "Entrada duplicada" in result:
                    # Já existe: atualiza (UPSERT lógico)
                    update_sql = """
                        UPDATE DizimoLancamentos
                           SET valor_dizimo = ?, valor_oferta = ?, data_pagamento = ?, forma_pagamento = ?, observacoes = ?, atualizado_em = SYSUTCDATETIME()
                         WHERE membro_id = ? AND ano = YEAR(?) AND mes = MONTH(?)
                    """
                    up_params = (float(valor_dizimo), float(valor_oferta), data_pagamento, forma_pagamento, (observacoes if observacoes.strip() else None),
                                 int(membro_escolhido), competencia, competencia)
                    ok = execute_query(update_sql, up_params)
                    if ok is True:
                        st.success("Lançamento atualizado (competência já existia).")
                    else:
                        st.error(f"Falha ao atualizar: {ok}")
                else:
                    st.error(f"Falha ao salvar: {result}")

    # ========= TAB 2: Painel anual (estilo planilha) =========
    with tab2:
        st.subheader("Visão anual por membro (meses em colunas)")

        ano_sel = st.number_input("Ano", min_value=1900, max_value=2100, value=datetime.date.today().year, step=1)
        # Busca totais mensais por membro
        query = """
            SELECT l.membro_id, m.nome,
                   l.ano, l.mes,
                   SUM(l.valor_dizimo)   AS total_dizimo,
                   SUM(l.valor_oferta)   AS total_oferta
              FROM DizimoLancamentos l
              JOIN Membros m ON m.id = l.membro_id
             WHERE l.ano = ?
             GROUP BY l.membro_id, m.nome, l.ano, l.mes
        """
        df = read_records(query, params=(int(ano_sel),))

        if df.empty:
            st.info("Sem lançamentos para este ano.")
        else:
            # Pivot para ficar igual à planilha de dizimistas (membros x 12 meses)
            pvt_diz = df.pivot_table(index=["membro_id","nome"], columns="mes", values="total_dizimo", aggfunc="sum", fill_value=0.0)
            pvt_oft = df.pivot_table(index=["membro_id","nome"], columns="mes", values="total_oferta", aggfunc="sum", fill_value=0.0)

            # Ordena colunas 1..12
            pvt_diz = pvt_diz.reindex(columns=range(1,13), fill_value=0.0)
            pvt_oft = pvt_oft.reindex(columns=range(1,13), fill_value=0.0)

            # Renomeia colunas para nomes de meses (abreviados PT-BR)
            meses = {i: datetime.date(2000,i,1).strftime("%b").capitalize() for i in range(1,13)}
            pvt_diz.rename(columns=meses, inplace=True)
            pvt_oft.rename(columns=meses, inplace=True)

            # Totais por linha
            pvt_diz["Total Dízimo"] = pvt_diz.sum(axis=1)
            pvt_oft["Total Ofertas"] = pvt_oft.sum(axis=1)

            # Junta em um único dataframe (colunas em blocos)
            painel = pd.concat(
                {
                    "Dízimo (R$)": pvt_diz,
                    "Ofertas (R$)": pvt_oft
                },
                axis=1
            )
            painel = painel.reset_index().rename(columns={"membro_id":"ID", "nome":"Membro"})

            st.dataframe(painel, use_container_width=True)

            # Exportar para Excel (achata MultiIndex)
            painel_excel = painel.copy()
            if isinstance(painel_excel.columns, pd.MultiIndex):
                painel_excel.columns = [
                    " - ".join([str(x) for x in col if x is not None and str(x) != ""])
                    for col in painel_excel.columns.to_flat_index()
                ]
            output = BytesIO()
            with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
                painel_excel.to_excel(writer, index=False, sheet_name=f"{ano_sel}")
            st.download_button(
                label="Baixar Excel do Painel",
                data=output.getvalue(),
                file_name=f"painel_dizimistas_{ano_sel}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )

            # KPIs simples
            total_ano = df["total_dizimo"].sum() + df["total_oferta"].sum()
            c1, c2, c3 = st.columns(3)
            c1.metric("Total Dízimos (ano)", f"R$ {df['total_dizimo'].sum():.2f}")
            c2.metric("Total Ofertas (ano)", f"R$ {df['total_oferta'].sum():.2f}")
            c3.metric("Total Geral (ano)",  f"R$ {total_ano:.2f}")

    # ========= TAB 3: Gerenciar lançamentos =========
    with tab3:
        st.subheader("Edição/Exclusão de lançamentos")
        # Filtros
        colf1, colf2, colf3 = st.columns(3)
        with colf1:
            ano_g = st.number_input("Ano", min_value=1900, max_value=2100, value=datetime.date.today().year, step=1, key="ano_g")
        with colf2:
            mes_g = st.selectbox("Mês", options=["Todos"] + list(range(1,13)), key="mes_g")
        with colf3:
            membro_g = st.text_input("Buscar por nome (contém)")

        base_q = """
            SELECT l.id, m.nome, l.ano, l.mes, l.valor_dizimo, l.valor_oferta, l.data_pagamento, l.forma_pagamento, l.observacoes
              FROM DizimoLancamentos l
              JOIN Membros m ON m.id = l.membro_id
             WHERE l.ano = ?
        """
        params = [int(ano_g)]
        if mes_g != "Todos":
            base_q += " AND l.mes = ?"
            params.append(int(mes_g))
        if membro_g.strip():
            base_q += " AND m.nome LIKE ?"
            params.append(f"%{membro_g.strip()}%")
        base_q += " ORDER BY m.nome, l.mes"

        lista = read_records(base_q, params=tuple(params))
        if lista.empty:
            st.info("Sem lançamentos no filtro.")
        else:
            # Exibir data_pagamento em PT-BR
            lista_br = df_to_br_display(lista, ["data_pagamento"])
            st.dataframe(lista_br, use_container_width=True)

            # Excluir
            with st.expander("Excluir lançamento"):
                id_del = st.selectbox("ID para excluir", options=lista["id"])
                if st.button("Confirmar exclusão"):
                    ok = execute_query("DELETE FROM DizimoLancamentos WHERE id = ?", (int(id_del),))
                    if ok is True:
                        st.success(f"Lançamento {id_del} excluído.")
                        st.rerun()
                    else:
                        st.error(f"Falha ao excluir: {ok}")

# -----------------------------------------------------------------------------
# Página 5 (exclusiva para adm-secretaria): Página para secretários
# -----------------------------------------------------------------------------

def page_secretaria():
    st.header("Página da Secretaria")
    st.write("Aqui você pode adicionar funcionalidades financeiras, por exemplo.")

# -----------------------------------------------------------------------------
# LÓGICA PRINCIPAL
# -----------------------------------------------------------------------------

def main():
    if not st.session_state["logged_in"]:
        login_screen()
        return
    logout_button()
    role = st.session_state["user_role"]
    pages = {
        "adm": {"Cadastro de Igreja": page_igreja, "Cadastro de Membros": page_membros, "Relatórios": page_relatorios},
        "adm-financeiro": {"Cadastro de Igreja": page_igreja, "Cadastro de Membros": page_membros, "Relatórios": page_relatorios, "Página Financeira": page_financeiro},
        "adm-secretaria": {"Cadastro de Igreja": page_igreja, "Cadastro de Membros": page_membros}
    }.get(role, {})
    if not pages:
        st.error("Usuário desconhecido. Verifique as credenciais.")
        return
    with st.sidebar:
        choice = st.selectbox("Selecione a Página", list(pages.keys()))
    pages[choice]()

if __name__ == "__main__":
    main()
